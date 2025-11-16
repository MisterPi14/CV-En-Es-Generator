import json
from pathlib import Path
from typing import Dict, Any
import weasyprint
from jinja2 import Environment, FileSystemLoader
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from rich.console import Console

console = Console()

class PDFGenerator:
    def __init__(self):
        self.templates_dir = Path("templates")
        self.html_templates_dir = self.templates_dir / "html"
        self.word_templates_dir = self.templates_dir / "word"
        
        # Setup Jinja2 environment for HTML templates
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.html_templates_dir))
        )
    
    def list_available_templates(self, template_type: str) -> list:
        """List available templates of specified type (html or word)"""
        if template_type == "html":
            templates_dir = self.html_templates_dir
            extension = ".html"
        elif template_type == "word":
            templates_dir = self.word_templates_dir
            extension = ".docx"
        else:
            return []
        
        if not templates_dir.exists():
            return []
        
        templates = []
        for template_file in templates_dir.glob(f"*{extension}"):
            templates.append(template_file.stem)
        
        return templates
    
    def generate_pdf_from_html(self, cv_data: Dict[Any, Any], template_name: str, output_path: str) -> bool:
        """Generate PDF using HTML template"""
        try:
            template_file = f"{template_name}.html"
            template = self.jinja_env.get_template(template_file)
            
            # Render HTML with CV data
            html_content = template.render(**cv_data)
            
            # Generate PDF
            weasyprint.HTML(string=html_content).write_pdf(output_path)
            return True
            
        except Exception as e:
            console.print(f"[red]Error generando PDF desde HTML: {e}[/red]")
            return False
    
    def generate_pdf_from_word(self, cv_data: Dict[Any, Any], template_name: str, output_path: str) -> bool:
        """Generate PDF using Word template"""
        try:
            template_path = self.word_templates_dir / f"{template_name}.docx"
            
            if not template_path.exists():
                console.print(f"[red]Plantilla Word no encontrada: {template_path}[/red]")
                return False
            
            # Load Word template
            doc = Document(str(template_path))
            
            # Replace placeholders in paragraphs
            self._replace_placeholders_in_doc(doc, cv_data)
            
            # Save temporary Word file
            temp_word_path = output_path.replace('.pdf', '_temp.docx')
            doc.save(temp_word_path)
            
            # Convert to PDF
            convert(temp_word_path, output_path)
            
            # Clean up temporary file
            Path(temp_word_path).unlink()
            
            return True
            
        except Exception as e:
            console.print(f"[red]Error generando PDF desde Word: {e}[/red]")
            return False
    
    def _replace_placeholders_in_doc(self, doc: Document, cv_data: Dict[Any, Any]):
        """Replace placeholders in Word document"""
        # Flatten CV data for easy replacement
        replacements = self._flatten_cv_data(cv_data)
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if f"{{{{{key}}}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        if f"{{{{{key}}}}}" in cell.text:
                            cell.text = cell.text.replace(f"{{{{{key}}}}}", str(value))
    
    def _flatten_cv_data(self, cv_data: Dict[Any, Any]) -> Dict[str, str]:
        """Flatten nested CV data for template replacement"""
        flat_data = {}
        
        # Personal info
        if 'personal_info' in cv_data:
            for key, value in cv_data['personal_info'].items():
                flat_data[f"personal_{key}"] = value or ""
        
        # Summary
        flat_data['summary'] = cv_data.get('summary', '')
        
        # Experience (first 3 jobs)
        if 'experience' in cv_data:
            for i, exp in enumerate(cv_data['experience'][:3]):
                flat_data[f"exp{i+1}_title"] = exp.get('title', '')
                flat_data[f"exp{i+1}_company"] = exp.get('company', '')
                flat_data[f"exp{i+1}_duration"] = exp.get('duration', '')
                flat_data[f"exp{i+1}_description"] = exp.get('description', '')
        
        # Education (first 2 entries)
        if 'education' in cv_data:
            for i, edu in enumerate(cv_data['education'][:2]):
                flat_data[f"edu{i+1}_degree"] = edu.get('degree', '')
                flat_data[f"edu{i+1}_institution"] = edu.get('institution', '')
                flat_data[f"edu{i+1}_year"] = edu.get('year', '')
        
        # Skills (join first 10)
        if 'skills' in cv_data:
            flat_data['skills'] = ', '.join(cv_data['skills'][:10])
        
        # Languages
        if 'languages' in cv_data:
            flat_data['languages'] = ', '.join(cv_data['languages'])
        
        # Certifications
        if 'certifications' in cv_data:
            flat_data['certifications'] = ', '.join(cv_data['certifications'])
        
        return flat_data