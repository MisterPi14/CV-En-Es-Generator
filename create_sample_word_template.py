from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sample_template():
    """Create a sample Word template for demonstration"""
    doc = Document()
    
    # Header
    header = doc.add_heading('{{personal_name}}', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact info
    contact = doc.add_paragraph('{{personal_email}} | {{personal_phone}} | {{personal_location}}')
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if doc.paragraphs:
        doc.add_paragraph()  # Add space
    
    # Summary
    doc.add_heading('PROFESSIONAL SUMMARY', level=1)
    doc.add_paragraph('{{summary}}')
    
    # Experience
    doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
    
    # Job 1
    doc.add_paragraph('{{exp1_title}} - {{exp1_company}}', style='Heading 2')
    doc.add_paragraph('{{exp1_duration}}').italic = True
    doc.add_paragraph('{{exp1_description}}')
    
    # Job 2
    doc.add_paragraph('{{exp2_title}} - {{exp2_company}}', style='Heading 2')
    doc.add_paragraph('{{exp2_duration}}').italic = True
    doc.add_paragraph('{{exp2_description}}')
    
    # Education
    doc.add_heading('EDUCATION', level=1)
    doc.add_paragraph('{{edu1_degree}} - {{edu1_institution}} ({{edu1_year}})')
    doc.add_paragraph('{{edu2_degree}} - {{edu2_institution}} ({{edu2_year}})')
    
    # Skills
    doc.add_heading('SKILLS', level=1)
    doc.add_paragraph('{{skills}}')
    
    # Languages
    doc.add_heading('LANGUAGES', level=1)
    doc.add_paragraph('{{languages}}')
    
    # Certifications
    doc.add_heading('CERTIFICATIONS', level=1)
    doc.add_paragraph('{{certifications}}')
    
    # Save template
    doc.save('templates/word/sample.docx')
    print("Sample Word template created: templates/word/sample.docx")

if __name__ == "__main__":
    create_sample_template()